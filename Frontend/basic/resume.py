from docx import Document

doc = Document()

doc.add_heading("NISHANT NAVIN JHA", level=1)
doc.add_paragraph("Phone: 9328894602\nEmail: nishantjha203@gmail.com\nLocation: Bhopal\nLinkedIn: linkedin.com/in/nishant-jha-3bbb6329b/\nGitHub: github.com/Nishant7621")

doc.add_heading("PROFILE SUMMARY", level=2)
doc.add_paragraph(
"Data Analyst skilled in Excel, SQL, Python, and BI tools with strong ability to design reports, automate workflows, and turn raw datasets into actionable insights. Experienced in forecasting, data modeling, dashboarding, and KPI tracking. Adept at building efficient pipelines, improving reporting speed, and supporting business decision-making."
)

doc.add_heading("SKILLS", level=2)
skills = [
"Advanced Excel, VBA Automation",
"Power BI, Tableau, DAX",
"MySQL, SQL Analytics",
"Python (Pandas, NumPy, Scikit-learn)",
"ETL, Power Query, Data Cleaning",
"HTML, CSS, JavaScript",
"MongoDB, PySpark, AWS"
]
for s in skills:
    doc.add_paragraph(f"• {s}")

doc.add_heading("PROJECTS", level=2)

doc.add_paragraph("Sales Performance Dashboard | SQL, Power BI, Python")
doc.add_paragraph(
"• Developed a complete analytics dashboard with forecasting, KPIs, and revenue insights.\n"
"• Built SQL pipelines and star-schema models; implemented 25+ DAX measures.\n"
"• Improved reporting speed by 90% and increased cross-selling opportunities by 5%."
)

doc.add_paragraph("Machine Learning Prediction Model | Python, Scikit-learn")
doc.add_paragraph(
"• Engineered a supervised learning model with feature engineering and hyperparameter tuning.\n"
"• Achieved 88% accuracy and automated end-to-end training pipeline.\n"
"• Visualized model insights using Matplotlib and Seaborn."
)

doc.add_paragraph("Data Cleaning & Analytics Project | SQL, Pandas")
doc.add_paragraph(
"• Cleaned and transformed 200K+ rows using SQL & Pandas.\n"
"• Designed KPI report with anomaly detection and trend analysis.\n"
"• Reduced manual reporting effort by 70% via automation."
)

doc.add_heading("EDUCATION", level=2)
doc.add_paragraph("B.Tech in Computer Science Engineering – Sage University Bhopal (2023–2027)\nCGPA: 8.6")
doc.add_paragraph("GSEB Intermediate")
doc.add_paragraph("CBSE 10th")

doc.add_heading("CERTIFICATIONS", level=2)
certs = [
"Introduction to MS Excel",
"Getting Started with AWS",
"Introduction to MongoDB",
"Power BI Workshop"
]
for c in certs:
    doc.add_paragraph(f"• {c}")

file_path = "/mnt/data/Nishant_Jha_Updated_Resume.docx"
doc.save(file_path)

file_path
